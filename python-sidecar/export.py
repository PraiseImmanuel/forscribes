"""Turns a stored transcript into a downloadable Markdown, DOCX, or TXT file."""
from pathlib import Path


def format_duration(seconds: float | None) -> str:
    if not seconds:
        return "unknown length"
    total = int(seconds)
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"{h}h {m}m"
    if m:
        return f"{m}m {s}s"
    return f"{s}s"


def export_transcript(transcript: dict, variant: str, fmt: str, dest_path: str) -> None:
    """variant: 'raw' or 'cleaned'. fmt: 'txt', 'md', or 'docx'."""
    text = transcript["raw_text"]
    if variant == "cleaned" and transcript.get("cleaned_text"):
        text = transcript["cleaned_text"]

    meta = f"{transcript['created_at']} · {format_duration(transcript['duration_seconds'])}"

    if fmt == "txt":
        Path(dest_path).write_text(text, encoding="utf-8")

    elif fmt == "md":
        content = f"# {transcript['title']}\n\n*{meta}*\n\n{text}\n"
        Path(dest_path).write_text(content, encoding="utf-8")

    elif fmt == "docx":
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(transcript["title"], level=1)
        meta_p = doc.add_paragraph(meta)
        meta_p.runs[0].italic = True
        meta_p.runs[0].font.size = Pt(9)

        for paragraph in text.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        doc.save(dest_path)

    else:
        raise ValueError(f"Unknown export format: {fmt}")
