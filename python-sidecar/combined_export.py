"""Combined export: turns a list of transcripts/documents into ONE
Markdown/DOCX file, or a ZIP of individual files - reachable from the
Library (multi-select), a Grouping session, or a Topic query's results.

Callers build a list of "blocks" rather than passing raw items, since the
right structure differs per source:
  - Library multi-select: a flat list of content blocks, no dividers.
  - A Grouping session: one divider per group (its name), followed by a
    content block per member - so group structure survives into the export.
  - A Topic query: a flat list of content blocks, each title already
    carries its rank/score (built by the caller), no dividers needed.

This module only assembles documents from already-resolved text; it doesn't
touch the database - routes_export.py does the resolving.
"""
import re
import zipfile
from pathlib import Path

from export import format_duration


def _safe_filename(title: str, max_len: int = 60) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]', "", title).strip()
    return (cleaned or "untitled")[:max_len]


def build_combined_markdown(
    blocks: list[dict],
    overall_title: str | None,
    include_metadata: bool,
) -> str:
    has_dividers = any(b["type"] == "divider" for b in blocks)
    content_level = "###" if has_dividers else "##"

    lines = []
    if overall_title:
        lines.append(f"# {overall_title}\n")

    for block in blocks:
        if block["type"] == "divider":
            lines.append(f"## {block['text']}\n")
        else:
            lines.append(f"{content_level} {block['title']}\n")
            if include_metadata and block.get("meta"):
                lines.append(f"*{block['meta']}*\n")
            lines.append(f"{block['text']}\n")

    return "\n".join(lines)


def build_combined_docx(
    blocks: list[dict],
    overall_title: str | None,
    include_metadata: bool,
    dest_path: str,
) -> None:
    from docx import Document
    from docx.shared import Pt

    has_dividers = any(b["type"] == "divider" for b in blocks)
    content_level = 3 if has_dividers else 2

    doc = Document()
    if overall_title:
        doc.add_heading(overall_title, level=1)

    for block in blocks:
        if block["type"] == "divider":
            doc.add_heading(block["text"], level=2)
        else:
            doc.add_heading(block["title"], level=content_level)
            if include_metadata and block.get("meta"):
                meta_p = doc.add_paragraph(block["meta"])
                meta_p.runs[0].italic = True
                meta_p.runs[0].font.size = Pt(9)
            for paragraph in block["text"].split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

    doc.save(dest_path)


def build_zip(
    blocks: list[dict],
    fmt: str,  # 'md' | 'docx' | 'txt'
    include_metadata: bool,
    dest_path: str,
) -> None:
    """One file per content block (dividers don't have their own content to
    export, so they're skipped - the group/rank context they carried is
    baked into each item's title instead)."""
    import tempfile

    content_blocks = [b for b in blocks if b["type"] == "content"]
    used_names: set[str] = set()

    with zipfile.ZipFile(dest_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, block in enumerate(content_blocks, start=1):
            base_name = _safe_filename(block["title"])
            name = base_name
            suffix = 2
            while name in used_names:
                name = f"{base_name} ({suffix})"
                suffix += 1
            used_names.add(name)

            with tempfile.NamedTemporaryFile(
                suffix=f".{fmt}", delete=False
            ) as tmp:
                tmp_path = tmp.name

            try:
                if fmt == "txt":
                    Path(tmp_path).write_text(block["text"], encoding="utf-8")
                elif fmt == "md":
                    content = f"# {block['title']}\n\n"
                    if include_metadata and block.get("meta"):
                        content += f"*{block['meta']}*\n\n"
                    content += f"{block['text']}\n"
                    Path(tmp_path).write_text(content, encoding="utf-8")
                elif fmt == "docx":
                    build_combined_docx([block], None, include_metadata, tmp_path)
                else:
                    raise ValueError(f"Unknown export format: {fmt}")

                zf.write(tmp_path, arcname=f"{i:02d} - {name}.{fmt}")
            finally:
                Path(tmp_path).unlink(missing_ok=True)
